import matplotlib.pyplot as plt
import numpy as np
from docx import Document
import re

# Load the .docx file
file_path = "C:\\Users\\ppp09\\Desktop\\單低.docx"  # Replace with the path to your .docx file
doc = Document(file_path)

# Collect all text from the document
all_text = []
for paragraph in doc.paragraphs:
    all_text.append(paragraph.text)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            all_text.append(cell.text)

# Initialize arrays for different headers
OLTP = []
OLAP = []
HTAP_70_30 = []
HTAP_50_50 = []

# Process the document line by line to populate corresponding arrays
current_header = None

for text in all_text:
    text = text.strip()
    if "OLTP:" in text:
        current_header = "OLTP"
    elif "OLAP:" in text:
        current_header = "OLAP"
    elif "Hybrid(70%OLTP, 30%OLAP):" in text:
        current_header = "HTAP_70_30"
    elif "Hybrid(50%OLTP, 50%OLAP):" in text:
        current_header = "HTAP_50_50"
    else:
        # Extract only the first percentage (assumed to be CPU %) for the current header
        match = re.search(r"\s([\d.]+)%", text)
        if match and current_header:
            value = float(match.group(1))
            if current_header == "OLTP":
                OLTP.append(value)
            elif current_header == "OLAP":
                OLAP.append(value)
            elif current_header == "HTAP_70_30":
                HTAP_70_30.append(value)
            elif current_header == "HTAP_50_50":
                HTAP_50_50.append(value)

# Generate x-axis values (0 to 60 seconds)
time = np.linspace(0, 60, max(len(OLTP), len(OLAP), len(HTAP_70_30), len(HTAP_50_50)))

plt.figure(figsize=(10, 6))

# Plot each dataset with corresponding color
if OLTP:
    plt.plot(np.linspace(0, 60, len(OLTP)), OLTP, label='OLTP', color='purple')
if OLAP:
    plt.plot(np.linspace(0, 60, len(OLAP)), OLAP, label='OLAP', color='red')
if HTAP_70_30:
    plt.plot(np.linspace(0, 60, len(HTAP_70_30)), HTAP_70_30, label='Hybrid (70% OLTP, 30% OLAP)', color="orange")
if HTAP_50_50:
    plt.plot(np.linspace(0, 60, len(HTAP_50_50)), HTAP_50_50, label='Hybrid (50% OLTP, 50% OLAP)', color='green')

# 修改圖例設置，水平排列
plt.legend(
    loc='upper center',  # 圖例放在圖表上方中央
    bbox_to_anchor=(0.5, 1.1),  # 調整圖例位置
    ncol=4,  # 設置為 4 列，橫向排列
    fontsize=12 # 圖例文字大小
)

# 計算平均值
average_values = {
    'OLTP': np.mean(OLTP) if OLTP else None,
    'OLAP': np.mean(OLAP) if OLAP else None,
    'Hybrid (70% OLTP, 30% OLAP)': np.mean(HTAP_70_30) if HTAP_70_30 else None,
    'Hybrid (50% OLTP, 50% OLAP)': np.mean(HTAP_50_50) if HTAP_50_50 else None,
}

# 格式化平均值文字
average_text = "\n".join(
    [f"{key} Average: {value:.2f}%" for key, value in average_values.items() if value is not None]
)

# 確保文字框顯示在圖表右側，並將邊框設置為與圖例一致的樣式
plt.gcf().text(
    0.65, 0.35, average_text, fontsize=12, va='center', ha='left',
    bbox=dict(boxstyle="round", edgecolor="#d0d0d0", facecolor="white", linewidth=1)
)
        
# Add labels, title, and legend
plt.xlabel('Time (seconds)')
plt.ylabel('CPU Usage (%)')
plt.title('CPU : multiple cores  Request per second : 2')
# 修改圖例設置，水平排列
plt.legend(
    loc='upper center',  # 圖例放在圖表上方中央
    bbox_to_anchor=(0.5, 1.1),  # 調整圖例位置
    ncol=4,  # 設置為 4 列，橫向排列
    fontsize=12 # 圖例文字大小
)
plt.grid(True)

# 調整整體佈局，確保圖表與圖例不重疊

plt.tight_layout(rect=[0, 0, 0.85, 1])  # 將圖表的寬度限制在 85%，留空間給圖例

# Display the graph
plt.show()
